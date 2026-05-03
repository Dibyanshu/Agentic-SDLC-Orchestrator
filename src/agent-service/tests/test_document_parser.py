import base64
import io
import unittest

from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DictionaryObject, DecodedStreamObject, NameObject

from app.context.document_parser import parse_source_content


class DocumentParserTests(unittest.TestCase):
    def test_txt_content_is_returned_unchanged(self) -> None:
        self.assertEqual("plain context", parse_source_content("txt", "plain context"))

    def test_docx_content_extracts_paragraphs_and_tables(self) -> None:
        document = Document()
        document.add_paragraph("DOCX paragraph context")
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "DOCX table context"
        buffer = io.BytesIO()
        document.save(buffer)

        parsed = parse_source_content("docx", base64.b64encode(buffer.getvalue()).decode("ascii"))

        self.assertIn("DOCX paragraph context", parsed)
        self.assertIn("DOCX table context", parsed)

    def test_pdf_content_extracts_text(self) -> None:
        pdf_bytes = _build_pdf_with_text("PDF RAG context")

        parsed = parse_source_content("pdf", base64.b64encode(pdf_bytes).decode("ascii"))

        self.assertIn("PDF RAG context", parsed)

    def test_binary_sources_require_valid_base64(self) -> None:
        with self.assertRaises(ValueError):
            parse_source_content("pdf", "not-base64")

    def test_unsupported_source_type_is_rejected(self) -> None:
        with self.assertRaises(ValueError):
            parse_source_content("md", "# Heading")


def _build_pdf_with_text(text: str) -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=300, height=144)

    stream = DecodedStreamObject()
    escaped_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream.set_data(f"BT /F1 18 Tf 72 72 Td ({escaped_text}) Tj ET".encode("ascii"))

    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font})}
    )
    page[NameObject("/Contents")] = stream

    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


if __name__ == "__main__":
    unittest.main()
