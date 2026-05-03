import base64
import io

from docx import Document
from pypdf import PdfReader


def parse_source_content(source_type: str, content: str) -> str:
    normalized_type = source_type.strip().lower()
    if normalized_type == "txt":
        return content
    if normalized_type == "pdf":
        return _parse_pdf(base64.b64decode(content, validate=True))
    if normalized_type == "docx":
        return _parse_docx(base64.b64decode(content, validate=True))

    raise ValueError("sourceType must be txt, pdf, or docx")


def _parse_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(page.strip() for page in pages if page.strip())


def _parse_docx(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    table_cells.append(text)

    return "\n".join([*paragraphs, *table_cells])
